from datetime import datetime
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

from open_webui.models.chats import ChatTitleMessagesForm

# Done: Checken ob der Pfad immer geht bzw dass man es so löst, dass er immer geht
# Done: Für Kandidatenprofil soll nur die ausgewählte Nachricht verwendet werden, bei der das Word Icon gedrückt wird,
# Done Nicht der ganze Chatverlauf
# Done: Name des Heruntergeladenen Dokuments anpassen: Schöne Namen, ohne emojis
# ? Vielleicht aufsplitten -> Einmal für gesamten Chatverlauf, einmal für eine explizite Nachricht
TEMPLATE_PATH = Path("backend/open_webui/static/templates/candidate_template.docx")


class WordGenerator:
    """
    Description:
    The `WordGenerator` class is designed to create Word (.docx) documents from chat messages.

    Attributes:
    - `form_data`: An instance of `ChatTitleMessagesForm` containing title and messages.
    """

    def __init__(self, form_data: ChatTitleMessagesForm):
        self.form_data = form_data

    def format_timestamp(self, timestamp: float) -> str:
        """Convert a UNIX timestamp to a formatted date string."""
        try:
            date_time = datetime.fromtimestamp(timestamp)
            return date_time.strftime("%Y-%m-%d, %H:%M:%S")
        except (ValueError, TypeError):
            return ""

    def generate_chat_word(self) -> bytes:
        """
        Generate a Word document from a template and fill it with chat messages.
        """
        try:
            doc = Document(TEMPLATE_PATH)

            # Add Title
            doc.add_heading(self.form_data.title, level=1)

            # Add Messages
            for msg in self.form_data.messages:
                # Add role, model, and timestamp
                role = msg.get("role", "User").title()
                model = f" ({msg.get('model')})" if msg.get("model") and role == "Assistant" else ""
                timestamp = msg.get("timestamp")
                date_str = self.format_timestamp(timestamp) if timestamp else ""
                heading = f"{role}{model} - {date_str}"
                doc.add_paragraph(heading)

                # Add content with bold formatting
                content = msg.get("content", "")
                self._add_formatted_text(doc, content)

                # Add separator
                doc.add_paragraph("-" * 40)

            # Align all paragraphs to the left
            for para in doc.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT

            # Save document to BytesIO
            word_io = BytesIO()
            doc.save(word_io)
            word_io.seek(0)

            return word_io.getvalue()
        except Exception as e:
            raise e

    def _add_formatted_text(self, doc, text: str):
        """
        Add text to the document with bold formatting for text enclosed in double asterisks.
        """
        paragraph = doc.add_paragraph()
        parts = text.split("**")
        for i, part in enumerate(parts):
            if i % 2 == 1:
                # Odd-indexed parts are inside double asterisks, so they should be bold
                run = paragraph.add_run(part)
                run.bold = True
            else:
                # Even-indexed parts are outside double asterisks, so they should be normal
                paragraph.add_run(part)
