from io import BytesIO
from pathlib import Path
from typing import Dict, Any, List
import re
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt
from open_webui.models.chats import ChatTitleMessagesForm

TEMPLATE_PATH = Path(__file__).parent.parent / "static" / "templates" / "Vorlage Kandidatenprofil_Kaan.docx"


class ChatWordGenerator:
    """
    Description:
    The `ChatWordGenerator` class is designed to create Word (.docx) documents by replacing placeholders
    in a template with chat message data.

    Attributes:
    - `form_data`: An instance of `ChatTitleMessagesForm` containing title and messages.
    """

    def __init__(self, form_data: ChatTitleMessagesForm):
        self.form_data = form_data

    def _extract_key_value_pairs(self, messages: List[Dict[str, Any]]) -> Dict[str, str]:
        """
        Extract key-value pairs from the messages where keys are capitalized (e.g., KANDIDAT).
        Handles multi-line values by preserving line breaks and stopping at the next key.

        Args:
        - `messages`: List of message dictionaries containing 'content'.

        Returns:
        - A dictionary mapping placeholder keys to their corresponding values.
        """
        key_value_pairs = {}
        for msg in messages:
            content = msg.get("content", "")
            lines = content.split("\n")
            i = 0
            while i < len(lines):
                line = lines[i].strip()
                if line.isupper():
                    key = line.strip(":").strip()
                    value_lines = []
                    j = i + 1
                    while j < len(lines):
                        next_line = lines[j].strip()
                        if next_line.isupper():
                            break
                        if next_line != "":
                            value_lines.append(next_line)
                        j += 1
                    value = "\n".join(value_lines)
                    key_value_pairs[key] = value
                    i = j
                else:
                    i += 1
        return key_value_pairs

    def _replace_placeholders_in_document(self, doc: Document, replacements: Dict[str, str]):
        """
        Replace placeholders in the Word document with the provided values.

        Args:
        - `doc`: The Document object to modify.
        - `replacements`: A dictionary mapping placeholder keys to their replacement values.
        """
        center_aligned_placeholders = {"UNTERNEHMEN", "POSITION"}
        bold_placeholders = {"KANDIDAT", "AKTUELLE POSITION"}
        background_placeholders = {
            "BERUFLICHER WERDEGANG",
            "TÄTIGKEITEN WÄHREND DES STUDIUMS",
            "AKADEMISCHE AUSBILDUNG",
            "BERUFSAUSBILDUNG",
            "ZIVILDIENST",
            "SCHULISCHE AUSBILDUNG",
        }

        for paragraph in doc.paragraphs:
            for key, value in replacements.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, value)
                    if key in background_placeholders:
                        if re.match(r"^\d", value):
                            paragraph_text = paragraph.text.strip()
                            lines = paragraph_text.split("\n")
                            paragraph.clear()
                            for line in lines:
                                if re.match(r"^\d", line):
                                    bold_run = paragraph.add_run(line + "\n")
                                    bold_run.bold = True
                                else:
                                    paragraph.add_run(line + "\n")
                    if key in center_aligned_placeholders:
                        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        for run in paragraph.runs:
                            run.font.name = "Arial"
                            run.font.size = Pt(12)
                    else:
                        paragraph.alignment = 0
                        paragraph.paragraph_format.line_spacing = 1.15
                        for run in paragraph.runs:
                            run.font.name = "Arial"
                            run.font.size = Pt(10)

        foto_placeholder = "{{FOTO}}"
        foto_path = Path(__file__).parent.parent / "static" / "templates" / "Bewerbungsfoto.png"
        foto_width = Cm(4)
        foto_height = Cm(3)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for key, value in replacements.items():
                        placeholder = f"{{{{{key}}}}}"
                        if placeholder in cell.text:
                            cell.text = cell.text.replace(placeholder, value)
                            for paragraph in cell.paragraphs:
                                paragraph.alignment = 0
                                if key in bold_placeholders:
                                    for run in paragraph.runs:
                                        if value in run.text:
                                            run.text = run.text.replace(value, "")
                                            new_run = paragraph.add_run(value)
                                            new_run.bold = True
                                            new_run.font.name = "Arial"
                                            new_run.font.size = Pt(10)
                                        else:
                                            run.font.name = "Arial"
                                            run.font.size = Pt(10)

                    if foto_placeholder in cell.text:
                        for paragraph in cell.paragraphs:
                            paragraph.clear()

                        cell.paragraphs[0].add_run().add_picture(str(foto_path), width=foto_width, height=foto_height)

    # Step 2: Remove any remaining placeholders (not replaced by values)
    def _remove_remaining_placeholders(self, doc):
        """remove any remaining placeholders in doc"""
        placeholder_pattern = re.compile(r"\{\{.*?\}\}")

        for paragraph in doc.paragraphs:
            matches = placeholder_pattern.findall(paragraph.text)
            if matches:
                for match in matches:
                    paragraph.text = paragraph.text.replace(match, "")

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    matches = placeholder_pattern.findall(cell.text)
                    if matches:
                        for match in matches:
                            cell.text = cell.text.replace(match, "")

    def generate_chat_word(self) -> bytes:
        """
        Generate a Word document from a template and replace placeholders with chat message data.

        Returns:
        - The generated Word document as bytes.
        """
        try:
            # Extract key-value pairs from the messages
            key_value_pairs = self._extract_key_value_pairs(self.form_data.messages)

            # Load the template document
            doc = Document(TEMPLATE_PATH)

            # Replace placeholders in the document
            self._replace_placeholders_in_document(doc, key_value_pairs)

            # Replace remaining placeholders with empty string
            self._remove_remaining_placeholders(doc)

            # Save the document to BytesIO
            word_io = BytesIO()
            doc.save(word_io)
            word_io.seek(0)

            return word_io.getvalue()
        except Exception as e:
            raise e
